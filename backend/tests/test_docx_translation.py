from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from app.engines.job_engine import _translate_docx


class FakeTranslationClient:
    target_language_code = "zh"
    target_language = "Simplified Chinese"

    def translate(self, text: str) -> str:
        return f"译文：{text}"


def test_docx_body_table_header_footer_are_translated_and_reopenable() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        source = root / "source.docx"
        output = root / "source_zh-CN.docx"

        document = Document()
        document.add_heading("Supplier Report", level=1)
        document.add_paragraph("This is a test paragraph.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Product"
        table.cell(0, 1).text = "Price"
        table.cell(1, 0).text = "Motor"
        table.cell(1, 1).text = "100 USD"
        document.sections[0].header.paragraphs[0].text = "Company Header"
        document.sections[0].footer.paragraphs[0].text = "Page Footer"
        document.save(source)

        result = _translate_docx(source, output, FakeTranslationClient(), None)
        reopened = Document(output)

        assert reopened.paragraphs[0].text == "译文：Supplier Report"
        assert reopened.paragraphs[1].text == "译文：This is a test paragraph."
        assert reopened.tables[0].cell(0, 0).text == "译文：Product"
        assert reopened.sections[0].header.paragraphs[0].text == "译文：Company Header"
        assert reopened.sections[0].footer.paragraphs[0].text == "译文：Page Footer"
        assert result["translated_items"] >= 5
        assert result["validation"]["sections"] == 1
        translated_run = reopened.paragraphs[0].runs[0]
        assert translated_run.font.name == "Microsoft YaHei"
        assert "Microsoft YaHei" in translated_run._element.xml
        assert result["validation"]["invalid_character_count"] == 0


def test_target_suffix_uses_language_code() -> None:
    from app.engines.job_engine import _target_suffix

    class Client:
        target_language_code = "zh"
        target_language = "Simplified Chinese"

    assert _target_suffix(Client()) == "zh-CN"


def test_split_runs_keep_structure_and_translate_all_visible_text() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        source = root / "complex.docx"
        output = root / "complex_zh-CN.docx"

        document = Document()
        paragraph = document.add_paragraph()
        first = paragraph.add_run("Supplier ")
        first.bold = True
        second = paragraph.add_run("Report")
        second.italic = True
        table = document.add_table(rows=1, cols=1)
        cell_p = table.cell(0, 0).paragraphs[0]
        cell_p.add_run("Servo ").bold = True
        cell_p.add_run("Motor")
        document.save(source)

        result = _translate_docx(source, output, FakeTranslationClient(), None)
        reopened = Document(output)

        assert reopened.paragraphs[0].text == "译文：Supplier Report"
        assert reopened.tables[0].cell(0, 0).text == "译文：Servo Motor"
        assert reopened.paragraphs[0].runs[0].bold is True
        assert result["translation_coverage"] == 100.0
        assert result["validation"]["text_blocks"] == result["source_text_blocks"]


def test_linked_headers_and_duplicate_text_are_counted_once_per_physical_paragraph() -> None:
    """Regression for V12.0.7 false 13->11 validation failures."""
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        source = root / "fidelity.docx"
        output = root / "fidelity_zh-CN.docx"

        document = Document()
        section = document.sections[0]
        section.header.paragraphs[0].text = "DEFAULT HEADER"
        # These objects may resolve to linked OOXML parts. The iterator must use
        # part-name + XPath, not Python object ids, to de-duplicate them safely.
        section.first_page_header.paragraphs[0].text = "FIRST HEADER"
        section.even_page_header.paragraphs[0].text = "EVEN HEADER"
        document.add_paragraph("Duplicate visible text")
        document.add_paragraph("Duplicate visible text")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Same cell text"
        table.cell(0, 1).text = "Same cell text"
        document.save(source)

        result = _translate_docx(source, output, FakeTranslationClient(), None)

        assert result["source_text_blocks"] == result["validation"]["text_blocks"]
        assert result["translation_coverage"] == 100.0
        reopened = Document(output)
        assert reopened.paragraphs[0].text == "译文：Duplicate visible text"
        assert reopened.paragraphs[1].text == "译文：Duplicate visible text"
        assert reopened.tables[0].cell(0, 0).text == "译文：Same cell text"
        assert reopened.tables[0].cell(0, 1).text == "译文：Same cell text"


def test_blank_provider_response_does_not_delete_source_text() -> None:
    class BlankClient(FakeTranslationClient):
        def translate(self, text: str) -> str:
            return ""

    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        source = root / "blank.docx"
        output = root / "blank_zh-CN.docx"
        document = Document()
        document.add_paragraph("Keep this content")
        document.save(source)

        result = _translate_docx(source, output, BlankClient(), None)
        reopened = Document(output)

        assert reopened.paragraphs[0].text == "Keep this content"
        assert result["skipped_items"] == 1
        assert result["validation"]["text_blocks"] == 1


def test_v1208_acceptance_suite_preserves_all_unique_text_blocks() -> None:
    suite = Path(__file__).resolve().parents[2] / "samples" / "acceptance"
    files = sorted(suite.glob("*.docx"))
    assert len(files) >= 12
    with TemporaryDirectory() as temp_dir:
        output_dir = Path(temp_dir)
        for source in files:
            output = output_dir / f"{source.stem}_zh-CN.docx"
            result = _translate_docx(source, output, FakeTranslationClient(), None)
            assert result["source_text_blocks"] == result["validation"]["text_blocks"], source.name
            assert result["validation"]["invalid_character_count"] == 0, source.name
            assert output.exists() and output.stat().st_size > 0, source.name

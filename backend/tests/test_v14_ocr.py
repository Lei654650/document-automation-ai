from pathlib import Path

from docx import Document

from app.engines.job_engine import _process_file


class EchoClient:
    target_language_code = "zh"
    target_language = "zh"
    source_language_code = "en"

    def translate(self, text: str) -> str:
        return f"译文：{text}"


def test_invoice_image_ocr_generates_non_empty_docx(tmp_path: Path):
    root = Path(__file__).resolve().parents[2]
    source = root / "samples" / "v13_acceptance" / "Invoice" / "01_Vietnam_VAT_Invoice.jpg"
    result = _process_file(source.name, str(source), tmp_path, EchoClient(), None, use_ocr=True)
    output = Path(result["path"])
    assert output.exists()
    assert output.stat().st_size > 0
    document = Document(output)
    text = "\n".join(p.text for p in document.paragraphs)
    assert "VAT" in text or "INVOICE" in text
    assert "译文" in text
    assert len(document.inline_shapes) >= 1


def test_ocr_delivery_rejects_empty_output(tmp_path: Path):
    # Regression guard: generated OCR Word must contain visible text.
    root = Path(__file__).resolve().parents[2]
    source = root / "samples" / "v13_acceptance" / "Image" / "02_Equipment_Nameplate.png"
    result = _process_file(source.name, str(source), tmp_path, EchoClient(), None, use_ocr=True)
    document = Document(result["path"])
    assert any(p.text.strip() for p in document.paragraphs)

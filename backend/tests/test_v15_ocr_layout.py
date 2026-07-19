from pathlib import Path

from docx import Document

from app.engines.job_engine import _build_editable_ocr_document, _detect_ocr_language


class FakeClient:
    target_language_code = "zh"


def test_invoice_reconstruction_is_editable_and_has_no_empty_cover():
    text = """增值税发票
供应商：Apex 自动化系统公司
客户：Nova Manufacturing Vietnam
发票号：INV-2026-0716 日期：2026-07-16
描述 数量 单价 金额
伺服电机 2 12,500,000 25,000,000
视觉相机 1 18,900,000 18,900,000
工程服务 1 35,000,000 35,000,000
总计：78,900,000 越南盾
付款：30 天内银行转账。"""
    source = Path(__file__).resolve().parents[2] / "samples" / "v14_acceptance" / "OCR" / "01_Vietnam_VAT_Invoice.jpg"
    document, blocks, template = _build_editable_ocr_document(source, text, text, FakeClient())
    assert template == "invoice"
    assert blocks >= 8
    assert document.paragraphs[0].text == "增值税发票"
    assert "OCR Document" not in "\n".join(p.text for p in document.paragraphs)
    assert len(document.tables) == 1
    assert len(document.tables[0].rows) == 4


def test_language_detection():
    assert _detect_ocr_language("VAT INVOICE Supplier") == "英文"
    assert _detect_ocr_language("增值税发票") == "中文"
    assert _detect_ocr_language("Hóa đơn VAT") == "越南语"
